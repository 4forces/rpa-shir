import rpa as r
import docx

# ========= Variables =========== #
d = docx.Document('D:\\Dropbox\\projects-own\\rpa-shir\\files\\cert-test2.docx')

print('Document object: ', d)
print('Document paragraphs: ', d.paragraphs)

# Selecting Name
name = d.paragraphs[4].text

# Selecting date line
d.paragraphs[9].text

# Selecting start date
start_date = d.paragraphs[9].runs[2].text

# Selecting end date
end_date = d.paragraphs[9].runs[5].text

